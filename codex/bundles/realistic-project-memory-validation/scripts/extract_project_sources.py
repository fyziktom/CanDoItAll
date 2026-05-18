from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import docx
import openpyxl
from pypdf import PdfReader


SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".xml", ".graphml"}
MAX_CELL_TEXT = 260
MAX_SHEET_ROWS_IN_MARKDOWN = 180
MAX_PDF_PAGES = 25


@dataclass(frozen=True)
class SourceProject:
    key: str
    title: str
    root: Path


def normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "-", value).strip("-").lower()


def read_xml_from_zip(zip_file: zipfile.ZipFile, name: str) -> ElementTree.Element | None:
    try:
        payload = zip_file.read(name)
    except KeyError:
        return None

    try:
        return ElementTree.fromstring(payload)
    except ElementTree.ParseError:
        return None


def xml_text(element: ElementTree.Element) -> str:
    parts: list[str] = []
    for value in element.itertext():
        clean = normalize_whitespace(unescape(value))
        if clean:
            parts.append(clean)
    return normalize_whitespace(" ".join(parts))


def extract_docx(path: Path) -> dict[str, Any]:
    document = docx.Document(str(path))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_whitespace(paragraph.text)
        if text:
            paragraphs.append(text)

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            values = [normalize_whitespace(cell.text) for cell in row.cells]
            if any(values):
                rows.append(values)
        if rows:
            tables.append({"index": table_index, "rows": rows})

    return {
        "kind": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_pdf(path: Path) -> dict[str, Any]:
    reader = PdfReader(str(path))
    pages: list[dict[str, Any]] = []
    for page_index, page in enumerate(reader.pages[:MAX_PDF_PAGES], start=1):
        text = normalize_whitespace(page.extract_text() or "")
        if text:
            pages.append({"page": page_index, "text": text})

    return {
        "kind": "pdf",
        "pageCount": len(reader.pages),
        "extractedPages": pages,
        "pageLimitApplied": len(reader.pages) > MAX_PDF_PAGES,
    }


def cell_display(value: Any) -> str | int | float | bool | None:
    if value is None:
        return None
    if isinstance(value, (int, float, bool)):
        return value
    return normalize_whitespace(str(value))[:MAX_CELL_TEXT]


def extract_workbook(path: Path) -> dict[str, Any]:
    workbook_values = openpyxl.load_workbook(path, data_only=True, read_only=False)
    workbook_formulas = openpyxl.load_workbook(path, data_only=False, read_only=False)

    sheets: list[dict[str, Any]] = []
    for sheet_name in workbook_values.sheetnames:
        sheet_values = workbook_values[sheet_name]
        sheet_formulas = workbook_formulas[sheet_name]

        cells: list[dict[str, Any]] = []
        non_empty_count = 0
        formula_count = 0
        for row in sheet_values.iter_rows():
            for value_cell in row:
                formula_cell = sheet_formulas[value_cell.coordinate]
                raw_formula = formula_cell.value
                value = cell_display(value_cell.value)
                formula = None
                if isinstance(raw_formula, str) and raw_formula.startswith("="):
                    formula = raw_formula[:MAX_CELL_TEXT]
                    formula_count += 1

                if value is None and formula is None:
                    continue

                non_empty_count += 1
                if len(cells) < 900:
                    cells.append(
                        {
                            "address": value_cell.coordinate,
                            "row": value_cell.row,
                            "column": value_cell.column,
                            "value": value,
                            "formula": formula,
                            "numberFormat": value_cell.number_format,
                        }
                    )

        merged_ranges = [str(merged_range) for merged_range in sheet_values.merged_cells.ranges]
        tables = [
            {
                "name": table.name,
                "displayName": table.displayName,
                "ref": table.ref,
            }
            for table in sheet_values.tables.values()
        ]

        sheets.append(
            {
                "name": sheet_name,
                "maxRow": sheet_values.max_row,
                "maxColumn": sheet_values.max_column,
                "nonEmptyCount": non_empty_count,
                "formulaCount": formula_count,
                "mergedRanges": merged_ranges,
                "tables": tables,
                "cells": cells,
                "truncated": non_empty_count > len(cells),
            }
        )

    return {
        "kind": "xlsx",
        "sheets": sheets,
    }


def extract_pptx(path: Path) -> dict[str, Any]:
    slides: list[dict[str, Any]] = []
    with zipfile.ZipFile(path) as package:
        slide_names = sorted(
            [
                name
                for name in package.namelist()
                if name.startswith("ppt/slides/slide") and name.endswith(".xml")
            ],
            key=lambda value: int(re.search(r"slide(\d+)\.xml", value).group(1)),
        )
        for slide_index, name in enumerate(slide_names, start=1):
            root = read_xml_from_zip(package, name)
            if root is None:
                continue
            texts: list[str] = []
            for element in root.iter():
                if element.tag.endswith("}t") and element.text:
                    clean = normalize_whitespace(element.text)
                    if clean:
                        texts.append(clean)
            slides.append({"slide": slide_index, "texts": texts})

    return {
        "kind": "pptx",
        "slides": slides,
    }


def extract_xmind_topic(topic: dict[str, Any]) -> dict[str, Any]:
    title = normalize_whitespace(str(topic.get("title", "")))
    notes_value = topic.get("notes")
    notes = ""
    if isinstance(notes_value, dict):
        notes = normalize_whitespace(str(notes_value.get("plain", {}).get("content", "")))
    children: list[dict[str, Any]] = []
    children_value = topic.get("children")
    if isinstance(children_value, dict):
        for attached in children_value.get("attached", []) or []:
            if isinstance(attached, dict):
                children.append(extract_xmind_topic(attached))
    return {"title": title, "notes": notes, "children": children}


def extract_xmind_zen(package: zipfile.ZipFile) -> dict[str, Any] | None:
    try:
        content = json.loads(package.read("content.json").decode("utf-8"))
    except Exception:
        return None

    sheets: list[dict[str, Any]] = []
    for sheet in content:
        root_topic = sheet.get("rootTopic")
        if not isinstance(root_topic, dict):
            continue
        sheets.append(
            {
                "title": normalize_whitespace(str(sheet.get("title", ""))),
                "root": extract_xmind_topic(root_topic),
            }
        )
    return {"format": "zen", "sheets": sheets}


def extract_xmind_legacy(package: zipfile.ZipFile) -> dict[str, Any] | None:
    root = read_xml_from_zip(package, "content.xml")
    if root is None:
        return None

    def local_name(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    def topic_from_xml(element: ElementTree.Element) -> dict[str, Any]:
        title = ""
        notes = ""
        children: list[dict[str, Any]] = []
        for child in element:
            name = local_name(child.tag)
            if name == "title":
                title = normalize_whitespace(xml_text(child))
            elif name == "notes":
                notes = normalize_whitespace(xml_text(child))
            elif name == "children":
                for topic_reference in child.iter():
                    if local_name(topic_reference.tag) == "topic":
                        children.append(topic_from_xml(topic_reference))
        return {"title": title, "notes": notes, "children": children}

    sheets: list[dict[str, Any]] = []
    for sheet in root.iter():
        if local_name(sheet.tag) != "sheet":
            continue
        sheet_title = ""
        root_topic: dict[str, Any] | None = None
        for child in sheet:
            name = local_name(child.tag)
            if name == "title":
                sheet_title = normalize_whitespace(xml_text(child))
            elif name == "topic":
                root_topic = topic_from_xml(child)
        if root_topic is not None:
            sheets.append({"title": sheet_title, "root": root_topic})
    return {"format": "legacy", "sheets": sheets}


def extract_xmind(path: Path) -> dict[str, Any]:
    with zipfile.ZipFile(path) as package:
        extracted = extract_xmind_zen(package) or extract_xmind_legacy(package)
        if extracted is None:
            return {"kind": "xmind", "format": "unknown", "entries": package.namelist()}
        return {"kind": "xmind", **extracted}


def extract_graphml(path: Path) -> dict[str, Any]:
    root = ElementTree.parse(path).getroot()
    nodes: list[dict[str, Any]] = []
    edges: list[dict[str, Any]] = []

    for element in root.iter():
        tag = element.tag.rsplit("}", 1)[-1]
        if tag == "node":
            label = ""
            for child in element.iter():
                if child.tag.rsplit("}", 1)[-1] in {"NodeLabel", "Label.Text"}:
                    label = normalize_whitespace(xml_text(child))
                    if label:
                        break
            nodes.append({"id": element.attrib.get("id", ""), "label": label})
        elif tag == "edge":
            label = ""
            for child in element.iter():
                if child.tag.rsplit("}", 1)[-1] in {"EdgeLabel", "Label.Text"}:
                    label = normalize_whitespace(xml_text(child))
                    if label:
                        break
            edges.append(
                {
                    "id": element.attrib.get("id", ""),
                    "source": element.attrib.get("source", ""),
                    "target": element.attrib.get("target", ""),
                    "label": label,
                }
            )

    return {"kind": "graphml", "nodes": nodes, "edges": edges}


def extract_image(path: Path) -> dict[str, Any]:
    metadata: dict[str, Any] = {"kind": "image", "sizeBytes": path.stat().st_size}
    try:
        from PIL import Image

        with Image.open(path) as image:
            metadata["width"] = image.width
            metadata["height"] = image.height
            metadata["format"] = image.format
    except Exception as error:
        metadata["dimensionError"] = str(error)
    return metadata


def extract_media(path: Path) -> dict[str, Any]:
    return {
        "kind": "media",
        "sizeBytes": path.stat().st_size,
        "extension": path.suffix.lower(),
    }


def extract_text_file(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="replace")
    return {"kind": path.suffix.lower().lstrip(".") or "text", "text": text}


def extract_file(path: Path) -> dict[str, Any]:
    extension = path.suffix.lower()
    if extension == ".docx":
        return extract_docx(path)
    if extension == ".pdf":
        return extract_pdf(path)
    if extension == ".xlsx":
        return extract_workbook(path)
    if extension == ".pptx":
        return extract_pptx(path)
    if extension == ".xmind":
        return extract_xmind(path)
    if extension == ".graphml":
        return extract_graphml(path)
    if extension in {".png", ".jpg", ".jpeg"}:
        return extract_image(path)
    if extension in {".mp4", ".mov", ".avi"}:
        return extract_media(path)
    if extension in SUPPORTED_TEXT_EXTENSIONS:
        return extract_text_file(path)
    return {
        "kind": "unknown",
        "sizeBytes": path.stat().st_size,
        "extension": extension,
    }


def flatten_xmind_topic(topic: dict[str, Any], depth: int = 0) -> list[str]:
    title = topic.get("title", "")
    notes = topic.get("notes", "")
    line = f"{'  ' * depth}- {title}"
    if notes:
        line += f": {notes}"
    lines = [line]
    for child in topic.get("children", []) or []:
        lines.extend(flatten_xmind_topic(child, depth + 1))
    return lines


def workbook_markdown(extracted: dict[str, Any]) -> list[str]:
    lines: list[str] = []
    for sheet in extracted.get("sheets", []):
        lines.append(f"### Sheet: {sheet['name']}")
        lines.append(
            f"- Used range: {sheet['maxRow']} rows x {sheet['maxColumn']} columns; non-empty cells: {sheet['nonEmptyCount']}; formulas: {sheet['formulaCount']}."
        )
        if sheet.get("tables"):
            lines.append(f"- Tables: {json.dumps(sheet['tables'], ensure_ascii=False)}")
        if sheet.get("mergedRanges"):
            lines.append(f"- Merged ranges: {', '.join(sheet['mergedRanges'][:20])}")
        cells = sheet.get("cells", [])
        if not cells:
            lines.append("- No non-empty cells extracted.")
            lines.append("")
            continue
        lines.append("")
        lines.append("| Cell | Value | Formula | Format |")
        lines.append("| --- | --- | --- | --- |")
        for cell in cells[:MAX_SHEET_ROWS_IN_MARKDOWN]:
            value = "" if cell.get("value") is None else str(cell.get("value")).replace("|", "\\|")
            formula = "" if cell.get("formula") is None else str(cell.get("formula")).replace("|", "\\|")
            number_format = str(cell.get("numberFormat", "")).replace("|", "\\|")
            lines.append(f"| {cell['address']} | {value} | {formula} | {number_format} |")
        if sheet.get("truncated"):
            lines.append(f"| ... | Truncated after {MAX_SHEET_ROWS_IN_MARKDOWN} markdown rows; full extraction JSON has {len(cells)} sampled cells. |  |  |")
        lines.append("")
    return lines


def markdown_for_file(project: SourceProject, file_path: Path, extracted: dict[str, Any]) -> str:
    relative = file_path.relative_to(project.root)
    lines = [
        f"# Extracted Source: {relative}",
        "",
        f"- Source path: `{file_path}`",
        f"- Source kind: `{extracted.get('kind', 'unknown')}`",
        f"- Project: `{project.title}`",
        "",
    ]

    kind = extracted.get("kind")
    if kind == "docx":
        lines.append("## Paragraphs")
        for paragraph in extracted.get("paragraphs", []):
            lines.append(f"- {paragraph}")
        for table in extracted.get("tables", []):
            lines.append("")
            lines.append(f"## Table {table['index']}")
            for row in table.get("rows", [])[:80]:
                lines.append("- " + " | ".join(row))
    elif kind == "pdf":
        lines.append(f"## PDF Pages")
        lines.append(f"- Page count: {extracted.get('pageCount')}")
        if extracted.get("pageLimitApplied"):
            lines.append(f"- Extracted first {MAX_PDF_PAGES} pages only.")
        for page in extracted.get("extractedPages", []):
            lines.append("")
            lines.append(f"### Page {page['page']}")
            lines.append(page["text"])
    elif kind == "xlsx":
        lines.extend(workbook_markdown(extracted))
    elif kind == "pptx":
        for slide in extracted.get("slides", []):
            lines.append(f"## Slide {slide['slide']}")
            for text in slide.get("texts", []):
                lines.append(f"- {text}")
            lines.append("")
    elif kind == "xmind":
        for sheet in extracted.get("sheets", []):
            lines.append(f"## Sheet: {sheet.get('title', '')}")
            root = sheet.get("root")
            if root:
                lines.extend(flatten_xmind_topic(root))
            lines.append("")
    elif kind == "graphml":
        lines.append("## Nodes")
        for node in extracted.get("nodes", []):
            lines.append(f"- `{node.get('id')}` {node.get('label')}")
        lines.append("")
        lines.append("## Edges")
        for edge in extracted.get("edges", []):
            label = f" ({edge.get('label')})" if edge.get("label") else ""
            lines.append(f"- `{edge.get('source')}` -> `{edge.get('target')}`{label}")
    elif kind == "image":
        lines.append("## Image Metadata")
        for key, value in extracted.items():
            if key != "kind":
                lines.append(f"- {key}: {value}")
    elif kind == "media":
        lines.append("## Media Metadata")
        for key, value in extracted.items():
            if key != "kind":
                lines.append(f"- {key}: {value}")
    elif "text" in extracted:
        lines.append("## Text")
        lines.append(extracted["text"])
    else:
        lines.append("## Metadata")
        lines.append(json.dumps(extracted, ensure_ascii=False, indent=2))

    return "\n".join(lines).rstrip() + "\n"


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_project(project: SourceProject, output_root: Path) -> dict[str, Any]:
    project_output = output_root / project.key
    raw_output = project_output / "raw"
    raw_output.mkdir(parents=True, exist_ok=True)

    files: list[dict[str, Any]] = []
    for file_path in sorted(project.root.rglob("*")):
        if not file_path.is_file():
            continue

        relative = file_path.relative_to(project.root)
        output_stem = safe_name(str(relative))
        markdown_path = raw_output / f"{output_stem}.md"
        json_path = raw_output / f"{output_stem}.json"

        try:
            extracted = extract_file(file_path)
            error = None
        except Exception as extract_error:
            extracted = {
                "kind": "error",
                "error": f"{type(extract_error).__name__}: {extract_error}",
            }
            error = extracted["error"]

        write_json(json_path, extracted)
        markdown_path.write_text(markdown_for_file(project, file_path, extracted), encoding="utf-8")
        files.append(
            {
                "path": str(file_path),
                "relativePath": str(relative),
                "kind": extracted.get("kind", "unknown"),
                "sizeBytes": file_path.stat().st_size,
                "markdown": str(markdown_path),
                "json": str(json_path),
                "error": error,
            }
        )

    index = {
        "key": project.key,
        "title": project.title,
        "root": str(project.root),
        "files": files,
    }
    write_json(project_output / "source-index.json", index)
    return index


def parse_arguments() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-root", required=True)
    parser.add_argument("--output-root", required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_arguments()
    input_root = Path(args.input_root).resolve()
    output_root = Path(args.output_root).resolve()

    projects = [
        SourceProject(
            key="ai-tap",
            title="AI Tap Intelligent Water Faucet",
            root=input_root / "AI kohoutek",
        ),
        SourceProject(
            key="curacao-glass-recycle",
            title="Curacao Glass Recycling and Foam Glass Plant",
            root=input_root / "Glass_Recycle_Curacao_Master_Pack_Government_Submission_Checked_2026-04-06",
        ),
    ]

    missing = [str(project.root) for project in projects if not project.root.exists()]
    if missing:
        print("Missing source project roots:", file=sys.stderr)
        for path in missing:
            print(f"- {path}", file=sys.stderr)
        return 2

    output_root.mkdir(parents=True, exist_ok=True)
    indexes = [extract_project(project, output_root) for project in projects]
    write_json(output_root / "source-index.json", {"projects": indexes})
    print(json.dumps({"outputRoot": str(output_root), "projects": indexes}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
